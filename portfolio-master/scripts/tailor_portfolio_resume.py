from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document


SOURCE = Path(r"C:\Users\ADMIN\Downloads\Francis_Gbohunmi_CV.docx")
OUTPUT = Path(__file__).resolve().parents[1] / "Francis_Gbohunmi_Portfolio_Resume.docx"


def replace_runs(paragraph, parts: list[tuple[str, str]]) -> None:
    templates = {
        "plain": paragraph.runs[0] if paragraph.runs else None,
        "bold": next((run for run in paragraph.runs if run.bold), None),
        "italic": next((run for run in paragraph.runs if run.italic), None),
    }
    fallback = templates["plain"]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)

    for text, kind in parts:
        run = paragraph.add_run(text)
        template = templates.get(kind) or fallback
        if template is not None and template._r.rPr is not None:
            run._r.insert(0, copy.deepcopy(template._r.rPr))
        if kind == "bold":
            run.bold = True
            run.italic = False
        elif kind == "italic":
            run.bold = False
            run.italic = True
        else:
            run.bold = False
            run.italic = False


def set_plain(paragraph, text: str) -> None:
    replace_runs(paragraph, [(text, "plain")])


def set_label(paragraph, label: str, text: str) -> None:
    replace_runs(paragraph, [(label, "bold"), (text, "plain")])


def set_role(paragraph, role: str, company: str, date: str) -> None:
    replace_runs(
        paragraph,
        [(role, "bold"), (f" | {company}", "plain"), (f"\t{date}", "bold")],
    )


def set_project(paragraph, name: str, stack: str) -> None:
    replace_runs(paragraph, [(name, "bold"), (f" | {stack}", "italic")])


def set_bullet(paragraph, parts: list[tuple[str, str]] | str) -> None:
    if isinstance(parts, str):
        parts = [(parts, "plain")]
    replace_runs(paragraph, parts)


def main() -> None:
    shutil.copyfile(SOURCE, OUTPUT)
    document = Document(OUTPUT)
    paragraphs = document.paragraphs

    set_plain(
        paragraphs[1],
        "Backend Software Engineer | C# / .NET | APIs | Real-Time & Offline Systems",
    )

    set_plain(
        paragraphs[4],
        "Backend-focused Software Engineer building reliable APIs and product systems "
        "with C#, ASP.NET Core, .NET 8/10, PostgreSQL, SQL Server, and SignalR. "
        "Experience spans fintech, enterprise ERP, edtech, music intelligence, personal "
        "safety, private social products, and farm operations. Skilled in Clean "
        "Architecture, CQRS, authentication, idempotent workflows, background jobs, "
        "offline sync, and API testing. I design around durable state, explicit failure "
        "handling, and clear contracts across web, mobile, and backend teams.",
    )

    set_label(
        paragraphs[6],
        "Languages & Backend: ",
        "C#, ASP.NET Core, .NET 8/10, Entity Framework Core, Python, TypeScript, SQL, LINQ",
    )
    set_label(
        paragraphs[7],
        "Architecture & APIs: ",
        "Clean Architecture, Modular Monoliths, CQRS (MediatR), REST APIs, JWT, SignalR, Background Jobs",
    )
    set_label(
        paragraphs[8],
        "Data & Infrastructure: ",
        "PostgreSQL, SQL Server, SQLite, Redis, Docker, Azure, AWS (EC2, S3, RDS, Lambda)",
    )
    set_label(
        paragraphs[9],
        "Product Engineering: ",
        "Offline-First Sync, Idempotency, Real-Time Systems, RBAC, Push Notifications",
    )
    set_label(
        paragraphs[10],
        "Tools & Quality: ",
        "Git, GitHub, Visual Studio, VS Code, xUnit, pytest, API Testing, Manual and Regression Testing",
    )

    set_role(paragraphs[12], "Backend Developer / QA Tester", "Model Carbon Ltd", "Dec 2025 - Jun 2026")
    set_bullet(
        paragraphs[13],
        [
            ("Built and validated backend services and REST APIs for ", "plain"),
            ("Converge ERP", "bold"),
            (", developed for Heirs Technologies under the UBA Group, using established modular .NET architecture and integration contracts.", "plain"),
        ],
    )
    set_bullet(
        paragraphs[14],
        "Debugged cross-service workflows, reviewed code, tracked regressions, and performed manual QA with frontend, AI, and QA teams ahead of releases.",
    )
    set_bullet(
        paragraphs[15],
        "Supported delivery from requirements and API documentation through integration, testing, and deployment readiness.",
    )

    set_role(paragraphs[16], "Backend Developer", "Kuda MFB", "Jun 2025 - Sep 2025")
    set_bullet(
        paragraphs[17],
        [
            ("Delivered backend workflows for ", "plain"),
            ("premium subscriptions, referral rewards, and loyalty", "bold"),
            (" using ASP.NET Core and SQL Server, with reliable transaction rules and account state.", "plain"),
        ],
    )
    set_bullet(
        paragraphs[18],
        "Collaborated with senior engineers to translate business goals into API contracts and launch engagement features for digital banking customers.",
    )
    set_bullet(
        paragraphs[19],
        "Implemented and tested eligibility, duplicate-action, reward issuance, and state-transition edge cases to improve release confidence.",
    )

    set_role(paragraphs[20], ".NET Developer", "X3 Lab", "Jan 2025 - Jun 2025")
    set_bullet(
        paragraphs[21],
        "Engineered core ASP.NET Core API endpoints and business logic for a university-focused edtech platform.",
    )
    set_bullet(
        paragraphs[22],
        "Built course registration, result management, and academic progress workflows with validation and role-aware access to student records.",
    )
    set_bullet(
        paragraphs[23],
        "Translated product requirements into reliable backend services and coordinated API integration with the frontend team.",
    )

    set_project(paragraphs[25], "NoteFusion", ".NET 10, Python / ONNX, Expo, PostgreSQL")
    set_bullet(
        paragraphs[26],
        "Designed a three-tier music transcription system where Expo clients call a .NET API that orchestrates a private FastAPI and ONNX worker and persists job results.",
    )
    set_bullet(
        paragraphs[27],
        "Built pitch and timing post-processing, key-aware movable-do solfa mapping, confidence handling, editable key correction, and MusicXML and MIDI output.",
    )

    set_project(paragraphs[28], "Recchx", ".NET 8, React Native, SignalR, PostgreSQL")
    set_bullet(
        paragraphs[29],
        "Built a personal safety network with trusted circles, background location, safe-arrival timers, emergency contacts, and covert SOS alerts.",
    )
    set_bullet(
        paragraphs[30],
        "Designed SOS-first offline queues, bounded location snapshots, idempotent batch sync, Ghost Mode privacy, and durable REST state with SignalR updates.",
    )

    set_project(paragraphs[31], "2gether", ".NET 10, Next.js, SignalR, PostgreSQL")
    set_bullet(
        paragraphs[32],
        "Built an invite-only couple application with validated partner linking, messaging, cycle tracking, live presence, location, and synchronous games.",
    )
    set_bullet(
        paragraphs[33],
        "Scoped data from authenticated couple state and made the server own game timers, turns, scores, consent, and reconnect snapshots.",
    )

    set_project(paragraphs[34], "SBZ Farms", ".NET 10, Expo, CQRS, PostgreSQL, SQLite")
    set_bullet(
        paragraphs[35],
        "Built an operations platform for birds, eggs, feed, health, sales, payroll, alerts, audit history, and role-based access.",
    )
    set_bullet(
        paragraphs[36],
        "Implemented a SQLite offline source of truth, ordered outbox replay, token refresh, exponential backoff, dead-letter handling, notifications, and CSV backups.",
    )

    document.core_properties.title = "Francis Gbohunmi Portfolio Resume"
    document.core_properties.subject = "Backend Software Engineer"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
